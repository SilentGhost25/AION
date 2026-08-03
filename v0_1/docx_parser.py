"""
AION Module: DOCX Parser
Extracts clean text from .docx files using python-docx
"""
from pathlib import Path


def extract_docx_text(file_path: str) -> str:
    """Extract clean readable text from a DOCX file."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        from docx import Document
        doc = Document(str(path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        result = "\n\n".join(parts)
        print(f"[DOCX] Extracted {len(result.split())} words from {path.name}")
        return result

    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")
