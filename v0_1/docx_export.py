"""
VTU-Compliant Word Document (.docx) Exporter for AION Generated Papers.
Converts the unified paper data structure directly into a professional .docx document.
"""

import io
import re
from typing import Any, Dict, List, Optional
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn


def _set_cell_background(cell, hex_color: str):
    """Set background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tc_pr.append(shd)


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell margins in dxa (1 pt = 20 dxa)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


def _clean_latex_math_for_doc(text: str) -> str:
    """Convert LaTeX symbols/math tags to clean readable unicode text for Word."""
    if not isinstance(text, str):
        return str(text or "")
    
    # Strip unresolved math tags
    t = re.sub(r'\[MATH:[^\]]+\]', '', text)
    t = re.sub(r'\\text\{([^}]+)\}', r'\1', t)
    t = re.sub(r'\\sigma_\{condition\}\(Relation\)', r'σ_condition(Relation)', t)
    t = re.sub(r'\\sigma', 'σ', t)
    t = re.sub(r'\\pi', 'π', t)
    t = re.sub(r'\\times', '×', t)
    t = re.sub(r'\\leq', '≤', t)
    t = re.sub(r'\\geq', '≥', t)
    t = re.sub(r'\\neq', '≠', t)
    t = re.sub(r'\\rightarrow', '→', t)
    t = re.sub(r'\\infty', '∞', t)
    t = re.sub(r'\\sum', '∑', t)
    t = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', t)
    t = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', t)
    t = re.sub(r'[\$\\]', '', t)
    t = re.sub(r'\s{2,}', ' ', t).strip()
    return t


def generate_docx_from_paper(paper_data: Dict[str, Any]) -> io.BytesIO:
    """
    Generates a VTU formatted .docx document from a paper dictionary.
    Returns BytesIO object containing the docx binary.
    """
    doc = Document()

    # Configure Margins (0.75 in)
    sections = doc.sections
    for sec in sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    # Extract metadata safely
    config = paper_data.get("config") or {}
    subject_name = (
        config.get("subjectName") or 
        paper_data.get("subject") or 
        paper_data.get("subject_name") or 
        "Course Examination"
    )
    subject_code = config.get("subjectCode") or paper_data.get("subject_code") or "VTU2026"
    exam_type = (config.get("examType") or paper_data.get("exam_type") or "IAT1").upper()
    dept = config.get("department") or "Department of Computer Science & Engineering"
    semester = config.get("semester") or "VI Semester"
    max_marks = config.get("maxMarks") or (100 if "SEE" in exam_type else 50)
    duration = config.get("duration") or ("3 Hours" if "SEE" in exam_type else "90 Minutes")

    exam_title_map = {
        "IAT1": "INTERNAL ASSESSMENT TEST - I",
        "IAT2": "INTERNAL ASSESSMENT TEST - II",
        "IAT3": "INTERNAL ASSESSMENT TEST - III",
        "IA": "INTERNAL ASSESSMENT TEST",
        "MID": "MID-TERM EXAMINATION",
        "SEE": "SEMESTER END EXAMINATION",
    }
    exam_full_title = exam_title_map.get(exam_type, f"{exam_type} EXAMINATION")

    # --- 1. Institutional Header ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = title_p.add_run("VISVESVARAYA TECHNOLOGICAL UNIVERSITY / INSTITUTION")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.name = "Calibri"

    dept_p = doc.add_paragraph()
    dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept_p.paragraph_format.space_after = Pt(2)
    r2 = dept_p.add_run(f"{dept}\n{exam_full_title}")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"

    # --- 2. Exam Meta Information Grid ---
    info_table = doc.add_table(rows=3, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False

    meta_items = [
        [("Course Title:", True), (subject_name, False), ("Course Code:", True), (subject_code, False)],
        [("Semester:", True), (f"{semester}", False), ("Max. Marks:", True), (str(max_marks), False)],
        [("Duration:", True), (duration, False), ("Date / Time:", True), ("As Scheduled", False)],
    ]

    for row_idx, row_data in enumerate(meta_items):
        row = info_table.rows[row_idx]
        for col_idx, (text, is_bold) in enumerate(row_data):
            cell = row.cells[col_idx]
            _set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.bold = is_bold
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            if is_bold:
                _set_cell_background(cell, "F1F5F9")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # --- 3. Instructions Line ---
    inst_p = doc.add_paragraph()
    inst_p.paragraph_format.space_after = Pt(8)
    inst_run = inst_p.add_run("Note: Answer any FIVE full questions, choosing ONE full question from each module.")
    inst_run.bold = True
    inst_run.font.size = Pt(9.5)
    inst_run.font.name = "Calibri"
    inst_run.font.italic = True

    # --- 4. Main Question Paper Table ---
    # Col widths: Q.No (0.6 in), Sub (0.4 in), Question text (4.2 in), Marks (0.6 in), CO (0.5 in), RBT (0.5 in)
    col_widths = [Inches(0.6), Inches(0.4), Inches(4.3), Inches(0.6), Inches(0.5), Inches(0.5)]

    # Extract modules / questions hierarchy
    modules_data = []
    if "modules" in paper_data and isinstance(paper_data["modules"], list) and paper_data["modules"]:
        for m in paper_data["modules"]:
            modules_data.append(m)
    elif "questions" in paper_data and isinstance(paper_data["questions"], list):
        # Group flat questions by module
        by_module: Dict[int, List[Any]] = {}
        for q in paper_data["questions"]:
            mod_num = q.get("module") or q.get("sectionNumber") or 1
            by_module.setdefault(mod_num, []).append(q)
        for mod_num in sorted(by_module.keys()):
            modules_data.append({
                "module_index": mod_num,
                "module_title": f"Module {mod_num}",
                "questions": by_module[mod_num]
            })

    # Create master table
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    headers = ["Q.No", "Part", "Question Description", "Marks", "CO", "RBT"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = col_widths[i]
        _set_cell_background(hdr_cells[i], "E2E8F0")
        _set_cell_margins(hdr_cells[i], top=80, bottom=80, left=80, right=80)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 2 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.name = "Calibri"

    q_counter = 1

    for mod_idx, mod in enumerate(modules_data, start=1):
        m_title = mod.get("module_title") or mod.get("title") or f"Module {mod_idx}"
        m_idx = mod.get("module_index") or mod_idx

        # Module Banner Row (Merged across all 6 columns)
        mod_row = table.add_row()
        mod_cell = mod_row.cells[0]
        for c in mod_row.cells[1:]:
            mod_cell.merge(c)
        _set_cell_background(mod_cell, "CBD5E1")
        _set_cell_margins(mod_cell, top=60, bottom=60, left=100, right=100)
        mp = mod_cell.paragraphs[0]
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_after = Pt(0)
        mr = mp.add_run(f"MODULE - {m_idx}: {m_title.upper()}")
        mr.bold = True
        mr.font.size = Pt(10)
        mr.font.name = "Calibri"

        questions = mod.get("questions") or []
        for mq_idx, q in enumerate(questions):
            # If second question in module, insert "OR" row
            if mq_idx > 0 and mq_idx % 2 == 1:
                or_row = table.add_row()
                or_cell = or_row.cells[0]
                for c in or_row.cells[1:]:
                    or_cell.merge(c)
                _set_cell_background(or_cell, "F8FAFC")
                _set_cell_margins(or_cell, top=40, bottom=40, left=80, right=80)
                orp = or_cell.paragraphs[0]
                orp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                orp.paragraph_format.space_after = Pt(0)
                or_run = orp.add_run("OR")
                or_run.bold = True
                or_run.font.size = Pt(9.5)
                or_run.font.name = "Calibri"

            sub_qs = q.get("sub_questions") or q.get("subQuestions") or []
            curr_q_no = q.get("question_number") or q.get("questionNumber") or q.get("qNo") or q_counter

            if not sub_qs:
                # Single question without sub-questions
                q_text = _clean_latex_math_for_doc(q.get("text") or q.get("question_text") or "")
                q_marks = q.get("marks") or 10
                q_co = q.get("co") or f"CO{min(mod_idx, 5)}"
                q_rbt = q.get("bloom") or q.get("bloom_level") or q.get("rbt") or "L2"

                row = table.add_row()
                cells = row.cells
                for ci, cw in enumerate(col_widths):
                    cells[ci].width = cw
                    _set_cell_margins(cells[ci], top=60, bottom=60, left=60, right=60)

                cells[0].paragraphs[0].add_run(f"Q{curr_q_no}").bold = True
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[1].paragraphs[0].add_run("-").alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[2].paragraphs[0].add_run(q_text)
                cells[3].paragraphs[0].add_run(str(q_marks)).alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[4].paragraphs[0].add_run(str(q_co)).alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[5].paragraphs[0].add_run(str(q_rbt)).alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                letters = ["a", "b", "c", "d", "e"]
                for s_idx, sq in enumerate(sub_qs):
                    s_label = sq.get("label") or sq.get("sub_label") or (letters[s_idx] if s_idx < len(letters) else f"({s_idx+1})")
                    s_text = _clean_latex_math_for_doc(sq.get("text") or sq.get("question_text") or "")
                    s_marks = sq.get("marks") or (6 if s_idx == 0 else 4)
                    s_co = sq.get("co") or f"CO{min(mod_idx, 5)}"
                    s_rbt = sq.get("bloom") or sq.get("bloom_level") or sq.get("rbt") or "L2"

                    row = table.add_row()
                    cells = row.cells
                    for ci, cw in enumerate(col_widths):
                        cells[ci].width = cw
                        _set_cell_margins(cells[ci], top=50, bottom=50, left=60, right=60)

                    if s_idx == 0:
                        p0 = cells[0].paragraphs[0]
                        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p0.add_run(f"Q{curr_q_no}").bold = True
                    else:
                        cells[0].paragraphs[0].add_run("")

                    p1 = cells[1].paragraphs[0]
                    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p1.add_run(f"({s_label})").bold = True

                    cells[2].paragraphs[0].add_run(s_text)
                    
                    p3 = cells[3].paragraphs[0]
                    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p3.add_run(str(s_marks))

                    p4 = cells[4].paragraphs[0]
                    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p4.add_run(str(s_co))

                    p5 = cells[5].paragraphs[0]
                    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p5.add_run(str(s_rbt))

            q_counter += 1

    # --- 5. Course Outcomes Mapping Table ---
    cos = paper_data.get("courseOutcomes") or [
        "Understand fundamental concepts and theoretical foundations",
        "Apply analytical methods and problem-solving techniques",
        "Implement algorithms and system architectures",
        "Analyze performance tradeoffs and design alternatives",
        "Evaluate solution quality and system specifications",
    ]

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    co_heading = doc.add_paragraph()
    co_heading.paragraph_format.space_after = Pt(4)
    co_h_run = co_heading.add_run("Course Outcomes (COs) Mapping:")
    co_h_run.bold = True
    co_h_run.font.size = Pt(10.5)
    co_h_run.font.name = "Calibri"

    co_table = doc.add_table(rows=len(cos) + 1, cols=2)
    co_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    co_table.autofit = False

    co_col_widths = [Inches(1.0), Inches(6.0)]
    co_hdr = co_table.rows[0].cells
    co_hdr[0].width = co_col_widths[0]
    co_hdr[1].width = co_col_widths[1]
    _set_cell_background(co_hdr[0], "E2E8F0")
    _set_cell_background(co_hdr[1], "E2E8F0")
    co_hdr[0].paragraphs[0].add_run("CO Code").bold = True
    co_hdr[1].paragraphs[0].add_run("Course Outcome Description").bold = True

    for i, co_desc in enumerate(cos, start=1):
        row_cells = co_table.rows[i].cells
        row_cells[0].width = co_col_widths[0]
        row_cells[1].width = co_col_widths[1]
        _set_cell_margins(row_cells[0], top=40, bottom=40, left=60, right=60)
        _set_cell_margins(row_cells[1], top=40, bottom=40, left=60, right=60)
        row_cells[0].paragraphs[0].add_run(f"CO{i}").bold = True
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].paragraphs[0].add_run(str(co_desc))

    # --- 6. Percentage of CO Coverage and Syllabus Coverage Tables ---
    raw_qs = paper_data.get("questions") or []
    all_qs_docx = []
    for q_item in raw_qs:
        sub_items = q_item.get("subQuestions") or q_item.get("sub_questions")
        if sub_items and isinstance(sub_items, list):
            all_qs_docx.extend(sub_items)
        else:
            all_qs_docx.append(q_item)

    total_marks_computed = sum(float(q.get("marks") or 0) for q in all_qs_docx) or 1.0

    co_totals: Dict[str, float] = {}
    for q in all_qs_docx:
        co_key = str(q.get("co") or q.get("coMapping") or "CO1").strip().upper()
        co_totals[co_key] = co_totals.get(co_key, 0.0) + float(q.get("marks") or 0)

    co_cov = {
        f"co{n}": round((co_totals.get(f"CO{n}", 0.0) / total_marks_computed) * 100)
        for n in range(1, 6)
    }

    mod_totals: Dict[int, float] = {}
    for q in all_qs_docx:
        m_val = q.get("module") or q.get("moduleIndex") or q.get("module_index") or 1
        try:
            m_num = int(m_val)
        except Exception:
            m_num = 1
        mod_totals[m_num] = mod_totals.get(m_num, 0.0) + float(q.get("marks") or 0)

    syl_cov = {
        f"s{n}": round((mod_totals.get(n, 0.0) / total_marks_computed) * 100)
        for n in range(1, 6)
    }

    # Add CO Coverage Table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    co_cov_h = doc.add_paragraph()
    co_cov_h.paragraph_format.space_after = Pt(3)
    co_cov_run = co_cov_h.add_run("Percentage of CO Coverage:")
    co_cov_run.bold = True
    co_cov_run.font.size = Pt(10.0)
    co_cov_run.font.name = "Calibri"

    co_cov_table = doc.add_table(rows=2, cols=6)
    co_cov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    co_cov_table.autofit = False
    cov_col_widths = [Inches(2.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)]

    hdr_cov_0 = co_cov_table.rows[0].cells
    hdr_cov_0[0].paragraphs[0].add_run("Course Outcomes").bold = True
    for c_i in range(1, 6):
        hdr_cov_0[c_i].paragraphs[0].add_run(f"CO{c_i}").bold = True
        hdr_cov_0[c_i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_cov_1 = co_cov_table.rows[1].cells
    row_cov_1[0].paragraphs[0].add_run("Percentage").bold = True
    for c_i in range(1, 6):
        row_cov_1[c_i].paragraphs[0].add_run(f"{co_cov.get(f'co{c_i}', 0)}%")
        row_cov_1[c_i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, r in enumerate(co_cov_table.rows):
        for ci, c in enumerate(r.cells):
            c.width = cov_col_widths[ci]
            if row_idx == 0:
                _set_cell_background(c, "E2E8F0")
            _set_cell_margins(c, top=40, bottom=40, left=50, right=50)

    # Add Syllabus Coverage Table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    syl_cov_h = doc.add_paragraph()
    syl_cov_h.paragraph_format.space_after = Pt(3)
    syl_cov_run = syl_cov_h.add_run("Percentage of Syllabus Coverage:")
    syl_cov_run.bold = True
    syl_cov_run.font.size = Pt(10.0)
    syl_cov_run.font.name = "Calibri"

    syl_cov_table = doc.add_table(rows=2, cols=6)
    syl_cov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    syl_cov_table.autofit = False

    hdr_syl_0 = syl_cov_table.rows[0].cells
    hdr_syl_0[0].paragraphs[0].add_run("Modules Covered").bold = True
    for c_i in range(1, 6):
        hdr_syl_0[c_i].paragraphs[0].add_run(str(c_i)).bold = True
        hdr_syl_0[c_i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row_syl_1 = syl_cov_table.rows[1].cells
    row_syl_1[0].paragraphs[0].add_run("Percentage").bold = True
    for c_i in range(1, 6):
        row_syl_1[c_i].paragraphs[0].add_run(f"{syl_cov.get(f's{c_i}', 0)}%")
        row_syl_1[c_i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, r in enumerate(syl_cov_table.rows):
        for ci, c in enumerate(r.cells):
            c.width = cov_col_widths[ci]
            if row_idx == 0:
                _set_cell_background(c, "E2E8F0")
            _set_cell_margins(c, top=40, bottom=40, left=50, right=50)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
