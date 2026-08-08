"""
AION SHP Stage 1 — File Diagnostics
=====================================
Inspects uploaded files before extraction begins.
Determines the true file type, content characteristics,
and the optimal extraction pipeline.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from .error_knowledge import ErrorKnowledgeBase, Severity


@dataclass
class FileProfile:
    """Complete diagnostic profile of an uploaded file."""
    path:                 str
    true_type:            str
    size_bytes:           int
    size_mb:              float
    estimated_words:      int
    contains_math:        bool = False
    contains_images:      bool = False
    contains_tables:      bool = False
    contains_code:        bool = False
    language:             str  = "English"
    noise_ratio:          float = 0.0
    recommended_pipeline: str = "text_direct"
    warnings:             list[str] = field(default_factory=list)
    fatal:                str = ""


class FileDiagnostics:
    """
    Stage 1: Profile every uploaded file before extraction.
    Routes to the correct extraction pipeline.
    """

    MAGIC = {
        b"%PDF":        "pdf",
        b"PK\x03\x04": "docx",
        b"\xff\xd8\xff": "jpg",
        b"\x89PNG":     "png",
        b"GIF8":        "gif",
        b"\x1f\x8b":   "gz",
        b"Rar!":        "rar",
    }

    def diagnose(self, file_path: str, kb: ErrorKnowledgeBase) -> FileProfile:
        path = Path(file_path)
        if not path.exists():
            return FileProfile(
                path=file_path, true_type="missing",
                size_bytes=0, size_mb=0, estimated_words=0,
                fatal=f"File not found: {file_path}",
            )

        size_bytes = path.stat().st_size
        size_mb    = round(size_bytes / (1024**2), 2)
        true_type  = self._detect_type(path)

        profile = FileProfile(
            path            = file_path,
            true_type       = true_type,
            size_bytes      = size_bytes,
            size_mb         = size_mb,
            estimated_words = 0,
        )

        if size_mb > 50:
            profile.warnings.append(
                f"Large file ({size_mb}MB). Extraction may be slow."
            )
        if size_mb > 200:
            profile.warnings.append(
                "Very large file. Consider splitting into modules first."
            )

        if true_type == "pdf":
            self._diagnose_pdf(profile, path)
        elif true_type == "docx":
            self._diagnose_docx(profile, path)
        elif true_type == "txt":
            self._diagnose_text(profile, path)
        else:
            profile.warnings.append(
                f"Unsupported or unknown file type: {true_type}"
            )

        profile.recommended_pipeline = self._select_pipeline(profile)

        print(f"[SHP-S1] File: {path.name}")
        print(f"  Type     : {profile.true_type}")
        print(f"  Size     : {profile.size_mb}MB")
        print(f"  Pipeline : {profile.recommended_pipeline}")
        print(f"  Math     : {profile.contains_math}")
        print(f"  Images   : {profile.contains_images}")

        return profile

    def _detect_type(self, path: Path) -> str:
        try:
            with open(path, "rb") as f:
                header = f.read(8)
            for magic, ftype in self.MAGIC.items():
                if header.startswith(magic):
                    return ftype
        except Exception:
            pass

        ext = path.suffix.lower()
        return {
            ".pdf": "pdf", ".docx": "docx", ".doc": "doc",
            ".txt": "txt", ".md":   "txt",  ".tex": "txt",
            ".jpg": "jpg", ".png":  "png",  ".jpeg": "jpg",
        }.get(ext, "unknown")

    def _diagnose_pdf(self, profile: FileProfile, path: Path):
        try:
            import fitz
            doc   = fitz.open(str(path))
            pages = len(doc)
            text  = ""

            img_count   = 0
            table_count = 0

            for page in doc:
                text      += page.get_text()
                img_count += len(page.get_images())
                blocks = page.get_text("blocks")
                if len(blocks) > 10:
                    table_count += 1

            doc.close()

            profile.estimated_words  = len(text.split())
            profile.contains_images  = img_count > 0
            profile.contains_tables  = table_count > 3
            profile.contains_math    = self._has_math(text)
            profile.contains_code    = self._has_code(text)
            profile.noise_ratio      = self._noise_ratio(text)

            chars_per_page = len(text) / max(1, pages)
            if chars_per_page < 100:
                profile.warnings.append(
                    "Low text density. PDF may be scanned — OCR recommended."
                )

        except ImportError:
            profile.warnings.append("PyMuPDF not installed. Using text fallback.")
        except Exception as e:
            profile.warnings.append(f"PDF diagnostics failed: {e}")

    def _diagnose_docx(self, profile: FileProfile, path: Path):
        try:
            import docx
            doc = docx.Document(str(path))

            text = "\n".join(p.text for p in doc.paragraphs)
            profile.estimated_words = len(text.split())
            profile.contains_math   = self._has_math(text)
            profile.contains_tables = len(doc.tables) > 0
            profile.contains_images = False
            profile.noise_ratio     = self._noise_ratio(text)

        except ImportError:
            profile.warnings.append("python-docx not installed.")
        except Exception as e:
            profile.warnings.append(f"DOCX diagnostics failed: {e}")

    def _diagnose_text(self, profile: FileProfile, path: Path):
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
            profile.estimated_words = len(text.split())
            profile.contains_math   = self._has_math(text)
            profile.contains_code   = self._has_code(text)
            profile.noise_ratio     = self._noise_ratio(text)
        except Exception as e:
            profile.warnings.append(f"Text diagnostics failed: {e}")

    def _select_pipeline(self, profile: FileProfile) -> str:
        if profile.true_type == "txt":
            return "text_direct"
        if profile.true_type == "docx":
            return "docx_parser"
        if profile.true_type == "pdf":
            if profile.noise_ratio > 0.30:
                return "ocr"
            if profile.contains_images and profile.estimated_words < 500:
                return "ocr"
            if profile.contains_math:
                return "pymupdf_math"
            return "pymupdf"
        return "text_direct"

    def _has_math(self, text: str) -> bool:
        indicators = [
            r'[∫∑∏√∂∇±×÷≤≥≠≈]',
            r'\$[^$]+\$',
            r'\\(?:int|sum|frac|sqrt|lim)\b',
            r'\b(?:d/dx|dy/dx)\b',
            r'[A-Za-z]\^[{(]?\d',
        ]
        for pat in indicators:
            if re.search(pat, text):
                return True
        return False

    def _has_code(self, text: str) -> bool:
        indicators = [
            r'\bdef\s+\w+\s*\(',
            r'\bclass\s+\w+',
            r'\bfor\s+\w+\s+in\b',
            r'#include\s*<',
            r'public\s+static\s+void\s+main',
        ]
        for pat in indicators:
            if re.search(pat, text):
                return True
        return False

    def _noise_ratio(self, text: str) -> float:
        if not text:
            return 0.0
        noise = sum(
            1 for ch in text
            if ord(ch) < 32 or (127 <= ord(ch) <= 159)
        )
        return round(noise / max(1, len(text)), 4)
