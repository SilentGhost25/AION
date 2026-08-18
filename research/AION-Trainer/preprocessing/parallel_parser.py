"""
Parallel Parser for large textbooks.

Strategy:
    900 pages -> 16 workers -> Parallel Parsing -> Merge -> Knowledge Objects

Memory strategy:
    Never load full book. Process 50-page batches.
    Page Stream -> Worker Pool -> Knowledge Queue -> Merge
"""

import os
import json
import logging
import multiprocessing as mp
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field, asdict
from concurrent.futures import ProcessPoolExecutor, as_completed

logger = logging.getLogger("aion.preprocessing")


@dataclass
class KnowledgeObject:
    """Atomic unit of extracted knowledge."""
    object_id: str = ""
    kind: str = "text"           # text | heading | algorithm | equation | image | table
    content: str = ""
    page: int = 0
    source_file: str = ""
    subject_code: str = ""
    module_hint: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


def _parse_pdf_worker(args: tuple) -> List[Dict]:
    """Worker function for parsing a batch of PDF pages."""
    file_path, start_page, end_page, subject_code, module_hint = args

    try:
        import fitz
    except ImportError:
        logger.error("PyMuPDF not installed")
        return []

    results = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(start_page, min(end_page, len(doc))):
            page = doc[page_num]
            text = page.get_text("text")

            # Split into blocks
            blocks = [b.strip() for b in text.split("\n\n") if b.strip()]

            for block in blocks:
                kind = _classify_block(block)
                results.append({
                    "object_id": f"{Path(file_path).stem}_p{page_num}_{len(results)}",
                    "kind": kind,
                    "content": block,
                    "page": page_num + 1,
                    "source_file": file_path,
                    "subject_code": subject_code,
                    "module_hint": module_hint,
                    "metadata": {},
                })

        doc.close()
    except Exception as e:
        logger.error(f"Error parsing {file_path} pages {start_page}-{end_page}: {e}")

    return results


def _parse_docx_worker(args: tuple) -> List[Dict]:
    """Worker function for parsing DOCX files."""
    file_path, subject_code, module_hint = args

    try:
        import docx as python_docx
    except ImportError:
        logger.error("python-docx not installed")
        return []

    results = []
    try:
        doc = python_docx.Document(file_path)

        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            kind = "heading" if para.style.name.startswith("Heading") else "text"
            results.append({
                "object_id": f"{Path(file_path).stem}_p0_{len(results)}",
                "kind": kind,
                "content": para.text.strip(),
                "page": 0,
                "source_file": file_path,
                "subject_code": subject_code,
                "module_hint": module_hint,
                "metadata": {},
            })

        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            results.append({
                "object_id": f"{Path(file_path).stem}_p0_{len(results)}",
                "kind": "table",
                "content": str(rows),
                "page": 0,
                "source_file": file_path,
                "subject_code": subject_code,
                "module_hint": module_hint,
                "metadata": {},
            })

    except Exception as e:
        logger.error(f"Error parsing {file_path}: {e}")

    return results


def _classify_block(text: str) -> str:
    """Classify a text block by its content."""
    import re

    if re.match(r"^\s*(\d+\.\d+|\d+\)|Module\s*\d+)", text, re.IGNORECASE):
        return "heading"
    if re.search(r"\b(algorithm|procedure|pseudocode)\b", text, re.IGNORECASE):
        return "algorithm"
    if re.search(r"[=≈≤≥∑∫√]|\\frac|\\sum", text):
        return "equation"
    return "text"


class ParallelParser:
    """
    Parallel parser for large textbooks.

    Usage:
        parser = ParallelParser(num_workers=8, batch_size=50)
        knowledge_objects = parser.parse_all(["book1.pdf", "book2.pdf"])
    """

    def __init__(
        self,
        num_workers: int = 4,
        batch_size: int = 50,
        subject_code: str = "",
    ):
        self.num_workers = num_workers
        self.batch_size = batch_size
        self.subject_code = subject_code

    def parse_all(self, file_paths: List[str]) -> List[KnowledgeObject]:
        """Parse all files and return merged Knowledge Objects."""
        all_objects = []

        for file_path in file_paths:
            suffix = Path(file_path).suffix.lower()
            if suffix == ".pdf":
                objects = self._parse_pdf_parallel(file_path)
            elif suffix == ".docx":
                objects = self._parse_docx(file_path)
            else:
                logger.warning(f"Skipping unsupported file: {file_path}")
                continue

            all_objects.extend(objects)
            logger.info(f"  Parsed {file_path}: {len(objects)} objects")

        logger.info(f"Total Knowledge Objects: {len(all_objects)}")
        return all_objects

    def _parse_pdf_parallel(self, file_path: str) -> List[KnowledgeObject]:
        """Parse PDF using multiple workers."""
        try:
            import fitz
            doc = fitz.open(file_path)
            total_pages = len(doc)
            doc.close()
        except Exception as e:
            logger.error(f"Cannot open PDF {file_path}: {e}")
            return []

        # Create page batches
        batches = []
        for start in range(0, total_pages, self.batch_size):
            end = min(start + self.batch_size, total_pages)
            batches.append((file_path, start, end, self.subject_code, 0))

        logger.info(f"  Processing {file_path}: {total_pages} pages in {len(batches)} batches")

        all_results = []
        with ProcessPoolExecutor(max_workers=self.num_workers) as executor:
            futures = {executor.submit(_parse_pdf_worker, batch): batch for batch in batches}

            for future in as_completed(futures):
                try:
                    results = future.result()
                    all_results.extend(results)
                except Exception as e:
                    batch = futures[future]
                    logger.error(f"Batch failed {batch}: {e}")

        return [KnowledgeObject(**obj) for obj in all_results]

    def _parse_docx(self, file_path: str) -> List[KnowledgeObject]:
        """Parse DOCX file."""
        args = (file_path, self.subject_code, 0)
        results = _parse_docx_worker(args)
        return [KnowledgeObject(**obj) for obj in results]
